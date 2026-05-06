"""
Generador de documentos .docx para Evangelista & Co.
Toma datos de un engagement y reemplaza placeholders en los templates.
"""
from pathlib import Path
from docx import Document
from datetime import datetime
import re
import structlog

logger = structlog.get_logger()

TEMPLATES_DIR = Path(__file__).parent / "templates_docx"

# Mapeo de nombres de template
TEMPLATE_MAP = {
    "propuesta_foundation": "propuesta_foundation.docx",
    "propuesta_architecture": "propuesta_architecture.docx",
    "contrato_foundation": "contrato_foundation.docx",
    "contrato_architecture": "contrato_architecture.docx",
    "dictamen_forense": "reporte_dictamen.docx",
    "orden_servicio": "orden_servicio.docx",
    "expediente_operativo": "expediente_operativo.docx",
}


def generate_document(template_name: str, data: dict) -> Path:
    """
    Genera un documento .docx reemplazando placeholders con datos reales.
    """
    template_file = TEMPLATES_DIR / TEMPLATE_MAP.get(template_name, "")
    if not template_file.exists():
        raise FileNotFoundError(f"Template no encontrado: {template_name}")
    
    doc = Document(str(template_file))
    
    # Construir el mapa de reemplazos
    replacements = _build_replacements(data)
    
    # Reemplazar en párrafos
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, replacements)
    
    # Reemplazar en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)
    
    # Reemplazar en headers/footers
    for section in doc.sections:
        for header in [section.header, section.first_page_header]:
            if header:
                for paragraph in header.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)
        for footer in [section.footer, section.first_page_footer]:
            if footer:
                for paragraph in footer.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)
    
    # Guardar en directorio temporal
    output_dir = Path("/tmp/evangelista_docs")
    output_dir.mkdir(parents=True, exist_ok=True)
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_name = re.sub(r'[^\w\-]', '_', data.get('cliente_nombre', 'documento'))
    output_path = output_dir / f"{template_name}_{safe_name}_{timestamp}.docx"
    
    doc.save(str(output_path))
    logger.info("docx_generated", template=template_name, path=str(output_path))
    
    return output_path


def _build_replacements(data: dict) -> dict[str, str]:
    """Construye el mapa de placeholder → valor."""
    now = datetime.now()
    
    # Valores base
    base = {
        "[DD/MM/AAAA]": now.strftime("%d/%m/%Y"),
        "[DD]": str(now.day),
        "[MES]": _mes_nombre(now.month),
        "[AAAA]": str(now.year),
        "[FECHA INICIO]": data.get("fecha_inicio", now.strftime("%d/%m/%Y")),
        "[FECHA FIN]": data.get("fecha_fin", now.strftime("%d/%m/%Y")),
    }
    
    # Datos del cliente
    cliente = {
        "[NOMBRE DEL CLIENTE]": data.get("cliente_nombre", ""),
        "[RAZÓN SOCIAL DEL CLIENTE]": data.get("razon_social", data.get("cliente_nombre", "")),
        "[EMPRESA DEL CLIENTE]": data.get("cliente_nombre", ""),
        "[RFC DEL CLIENTE]": data.get("rfc_cliente", "[RFC]"),
        "[DOMICILIO DEL CLIENTE]": data.get("domicilio_cliente", "[DOMICILIO]"),
        "[SECTOR]": data.get("sector", ""),
        "[SECTOR --- ej. Manufactura Textil]": data.get("sector", ""),
        "[CIUDAD CLIENTE]": data.get("ciudad", "Puebla"),
        "[NOMBRE DEL DIRECTOR GENERAL O REPRESENTANTE]": data.get("contacto_nombre", ""),
        "[NOMBRE DEL REPRESENTANTE LEGAL]": data.get("contacto_nombre", ""),
        "[CARGO": data.get("contacto_cargo", "Director General"),
    }
    
    # Datos de Evangelista
    evangelista = {
        "[RAZÓN SOCIAL EVANGELISTA & CO.]": "Evangelista & Co.",
        "[RAZÓN SOCIAL OFICIAL]": "Evangelista & Co.",
        "[RAZÓN SOCIAL OFICIAL DE EVANGELISTA & CO.]": "Evangelista & Co.",
        "[RFC]": data.get("rfc_evangelista", "[RFC EVANGELISTA]"),
        "[RFC DEL SOCIO / FIRMA]": data.get("rfc_evangelista", "[RFC EVANGELISTA]"),
        "[DOMICILIO FISCAL COMPLETO]": data.get("domicilio_evangelista", "Puebla, Pue."),
        "[NOMBRE CEO EVANGELISTA]": data.get("ceo_nombre", ""),
        "[NOMBRE CTO EVANGELISTA]": data.get("cto_nombre", ""),
        "[NOMBRE CEO]": data.get("ceo_nombre", ""),
        "[NOMBRE DEL": data.get("ceo_nombre", ""),
        "[NOMBRE]": data.get("ceo_nombre", ""),
    }
    
    # Datos financieros
    gamma = data.get("factor_gamma", 0)
    setup_fee = data.get("setup_fee", 0)
    foundation_fee = data.get("foundation_fee", 0)
    total_impacto = data.get("total_impacto", 0)
    
    financiero = {
        "[VALOR Γ CALCULADO]": f"{gamma:.2f}" if gamma else "[Γ]",
        "[VALOR Γ]": f"{gamma:.2f}" if gamma else "[Γ]",
        "$[MONTO SETUP]": f"${setup_fee:,.0f}" if setup_fee else "$[MONTO]",
        "$[MONTO TOTAL]": f"${total_impacto:,.0f}" if total_impacto else "$[MONTO]",
        "$[MONTO]": f"${foundation_fee:,.0f}" if foundation_fee else "$[MONTO]",
        "[MONTO SETUP]": f"{setup_fee:,.0f}" if setup_fee else "[MONTO]",
        "[MONTO TOTAL]": f"{total_impacto:,.0f}" if total_impacto else "[MONTO]",
        "[X] semanas": f"{_semanas_por_gamma(gamma)} semanas" if gamma else "[X] semanas",
    }
    
    # Datos del nodo crítico
    nodo = {
        "[ÁREA]": data.get("nodo_critico", ""),
        "[ÁREA IDENTIFICADA EN SCOPING --- ej. Almacén / Inventarios / Ventas]": data.get("nodo_critico", ""),
        "[DESCRIPCIÓN BREVE DEL PROBLEMA OBSERVADO]": data.get("problema_preliminar", ""),
        "[ESTIMADO PRELIMINAR EN PESOS O PORCENTAJE, si aplica]": f"${total_impacto:,.0f} MXN" if total_impacto else "",
    }
    
    # Folios
    folio_num = data.get("folio_num", "001")
    folios = {
        "[EVA-F-001]": f"EVA-F-{folio_num}",
        "[EVA-A-001]": f"EVA-A-{folio_num}",
        "[EVA-D-001]": f"EVA-D-{folio_num}",
        "[EVA-C-F-001]": f"EVA-C-F-{folio_num}",
        "[EVA-C-A-001]": f"EVA-C-A-{folio_num}",
        "[EVA-OS-001]": f"EVA-OS-{folio_num}",
        "[EVA-EXP-001]": f"EVA-EXP-{folio_num}",
    }
    
    # Hallazgos (para dictamen)
    hallazgos = data.get("hallazgos", [])
    hallazgo_replacements = {}
    for i, h in enumerate(hallazgos[:4]):
        idx = i + 1
        hallazgo_replacements[f"[NOMBRE DEL HALLAZGO {idx}]"] = h.get("nombre", "")
        hallazgo_replacements[f"[NOMBRE DEL HALLAZGO {idx} --- ej."] = h.get("nombre", "")
    
    # Combinar todo
    all_replacements = {**base, **cliente, **evangelista, **financiero, **nodo, **folios, **hallazgo_replacements}
    return all_replacements


def _replace_in_paragraph(paragraph, replacements: dict[str, str]):
    """Reemplaza placeholders en un párrafo preservando el formato."""
    full_text = paragraph.text
    if not any(key in full_text for key in replacements):
        return
    
    new_text = full_text
    for placeholder, value in replacements.items():
        if placeholder in new_text:
            new_text = new_text.replace(placeholder, str(value))
    
    if new_text == full_text:
        return
    
    if paragraph.runs:
        first_run = paragraph.runs[0]
        first_run.text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = new_text


def _mes_nombre(mes: int) -> str:
    meses = ["enero", "febrero", "marzo", "abril", "mayo", "junio",
             "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]
    return meses[mes - 1] if 1 <= mes <= 12 else ""


def _semanas_por_gamma(gamma: float) -> str:
    if gamma <= 1.5: return "6-7"
    elif gamma <= 2.0: return "7-9"
    elif gamma <= 3.0: return "9-12"
    else: return "12+"
